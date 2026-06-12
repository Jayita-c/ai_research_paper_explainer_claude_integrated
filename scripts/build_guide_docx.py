"""Generate the project build guide as a .docx document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACCENT = RGBColor(0x2B, 0x57, 0x9A)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def body(text, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    doc.add_paragraph(text, style="List Number")


def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    for i, line in enumerate(text.strip("\n").split("\n")):
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        if i < len(text.strip("\n").split("\n")) - 1:
            p.add_run().add_break()
    return p


def prompt_box(label, text):
    lp = doc.add_paragraph()
    lr = lp.add_run(label)
    lr.bold = True
    lr.font.color.rgb = ACCENT
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = GREY


# ============================ TITLE ============================
title = doc.add_heading("AI Research Paper Explainer", level=0)
sub = doc.add_paragraph()
sr = sub.add_run("End-to-End Build Guide — v0 + GitHub + Claude Code")
sr.bold = True
sr.font.size = Pt(13)
sr.font.color.rgb = ACCENT
meta = doc.add_paragraph()
mr = meta.add_run(
    "Stack: Next.js 16 (App Router) · TypeScript · Vercel AI SDK · "
    "@ai-sdk/anthropic (Claude Opus 4.8) · unpdf · Tailwind / shadcn-ui"
)
mr.italic = True
mr.font.size = Pt(9.5)
mr.font.color.rgb = GREY

doc.add_paragraph()
body(
    "This guide documents the full workflow used to build the application: "
    "generating a UI with v0, wiring it to Claude through the Vercel AI SDK, "
    "adding features, and the setup/debugging steps required to actually run it. "
    "Steps marked NEW were added to complete the original outline."
)

# ============================ OVERVIEW ============================
heading("Workflow at a Glance", level=1)
for s in [
    "Step 1 — Generate the base UI with v0",
    "Step 2 — Connect to GitHub and clone the repo",
    "Step 3 — Environment & dependencies  (NEW)",
    "Step 4 — Implement the backend with Claude Code (Vercel AI SDK)",
    "Step 5 — Run & verify the app  (NEW)",
    "Step 6 — Fix the file-upload dropzone bug  (NEW)",
    "Step 7 — Add 'Explain Like Different Personas'",
    "Step 8 — Convert analysis generation to streaming",
    "Step 9 — Deployment  (NEW, optional)",
]:
    bullet(s)

# ============================ STEP 1 ============================
heading("Step 1 — Generate the Base UI with v0", level=1)
body("Open v0 (v0.app) and use the following prompt to scaffold the application.")
prompt_box(
    "v0 prompt:",
    "Build a modern AI-powered Research Paper Explainer.\n\n"
    "Features: Upload PDF paper; Chat with paper; Generate summary; "
    "Key Contributions section; Methodology section; Limitations section; "
    "Future Work section.\n\n"
    "Design: Modern SaaS dashboard; Clean academic style; Dark mode support; "
    "Responsive layout.\n\n"
    "Include: Upload area; Summary cards; Chat interface; Citation panel. "
    "Basic app version ready to use.",
)
body("Result: a working Next.js front end with all UI components, but with placeholder / non-functional backend logic.")

# ============================ STEP 2 ============================
heading("Step 2 — Connect to GitHub and Clone the Repo", level=1)
numbered("In v0, connect the project to GitHub and push it to a new repository.")
numbered("Clone the repository locally:")
code_block("git clone https://github.com/<your-username>/ai-research-explainer.git\ncd ai-research-explainer")

# ============================ STEP 3 (NEW) ============================
heading("Step 3 — Environment & Dependencies  (NEW)", level=1)
body(
    "This step is required before anything can run and was missing from the "
    "original outline. The app calls the Anthropic API, so it needs a key.",
)
numbered("Create an Anthropic API key at console.anthropic.com (Settings → API Keys).")
numbered("Create your local environment file from the template:")
code_block("cp .env.example .env.local")
numbered("Add your key to .env.local (it is git-ignored, so it will not be committed):")
code_block("ANTHROPIC_API_KEY=sk-ant-...")
numbered("Install dependencies (requires Node.js 20+ for Next.js 16):")
code_block("npm install")
body(
    "Note: the Claude Code step below adds two packages to package.json — "
    "@ai-sdk/anthropic (the official Anthropic provider for the Vercel AI SDK) "
    "and unpdf (serverless PDF text extraction). Re-run npm install if you "
    "pull changes that modify package.json.",
    italic=True,
)

# ============================ STEP 4 ============================
heading("Step 4 — Implement the Backend with Claude Code", level=1)
body("Open the cloned project in Claude Code and run the following prompt.")
prompt_box(
    "Claude Code prompt:",
    "This is a Next.js application generated by v0. I want users to upload a "
    "research paper PDF and receive: 1. Executive Summary 2. Key Contributions "
    "3. Methodology 4. Limitations 5. Future Work.\n\n"
    "Implement: PDF upload endpoint; PDF text extraction; Claude API integration; "
    "Structured output generation; Connect the generated UI to the backend; "
    "Error handling; Loading states. Use the Vercel AI SDK and TypeScript.",
)
body("What this produced:")
bullet("app/api/analyze/route.ts — multipart PDF upload endpoint (validates type, size, empty files).")
bullet("lib/pdf.ts — server-side text extraction with unpdf, detecting scanned/image-only PDFs.")
bullet("Structured output via the Vercel AI SDK (Zod schema) against model claude-opus-4-8.")
bullet("app/api/chat/route.ts — 'chat with the paper' endpoint using the same provider.")
bullet("Front end wired to send the file and render Summary, Key Contributions, Methodology, Limitations, Future Work, and Citations.")
bullet("Error handling (config/validation/extraction/rate-limit) and loading states.")
body(
    "Design choice: the official @ai-sdk/anthropic provider is used (driven by "
    "ANTHROPIC_API_KEY) so the app calls the Claude API directly, rather than "
    "routing through a gateway.",
    italic=True,
)

# ============================ STEP 5 (NEW) ============================
heading("Step 5 — Run & Verify the App  (NEW)", level=1)
numbered("Start the development server:")
code_block("npm run dev")
numbered("Open http://localhost:3000.")
numbered("Upload a text-based research-paper PDF (under 20 MB) and confirm the analysis sections populate.")
numbered("Optionally build for production to catch type/bundle issues:")
code_block("npm run build")

# ============================ STEP 6 (NEW) ============================
heading("Step 6 — Fix the File-Upload Dropzone Bug  (NEW)", level=1)
body(
    "On first run the upload control did nothing when clicked. Two real bugs "
    "in components/upload-area.tsx were found and fixed:",
)
bullet(
    "Click loop: the hidden file <input> was a child of the clickable <div>, "
    "so calling input.click() dispatched a click that bubbled back to the div's "
    "onClick and re-triggered .click() — the browser defends against this by "
    "refusing to open the dialog."
)
bullet(
    "Silent MIME rejection: the client rejected any PDF whose file.type was "
    "empty or non-standard (common on drag-drop and across OS/browser combos) "
    "with no feedback — stricter than the server."
)
body("Fix:")
bullet("Wrap the <input> in a native <label>, so clicking opens the dialog natively (no JS .click(), no bubbling loop, keyboard-accessible).")
bullet("Accept a PDF by extension or MIME type (accept=\"application/pdf,.pdf\" plus a .pdf name check) and show a toast on non-PDF files.")
bullet("Reset the input value on click so the same file can be re-selected after an error.")
body(
    "Lesson: a v0-generated dropzone using a programmatic .click() on a nested "
    "input is a common footgun; prefer a <label>-wrapped input.",
    italic=True,
)

# ============================ STEP 7 ============================
heading("Step 7 — Add 'Explain Like Different Personas'", level=1)
body("Add audience-specific explanation buttons.")
prompt_box(
    "Feature spec / Claude Code prompt:",
    "Add an 'Explain like a persona' feature. Buttons: Researcher, Product "
    "Manager, Executive, Student. When clicked, Claude receives a request such "
    "as 'Explain this paper for a product manager' and returns a tailored "
    "explanation.",
)
body("What this produced:")
bullet("app/api/explain/route.ts — a streaming endpoint with a distinct system prompt per persona (audience, focus, tone, jargon level).")
bullet("components/persona-explainer.tsx — the four persona buttons rendered in a new 'Explain' tab, streaming the response with loading and error states.")

# ============================ STEP 8 ============================
heading("Step 8 — Convert Analysis Generation to Streaming", level=1)
prompt_box(
    "Claude Code prompt:",
    "Convert all analysis generation to streaming responses using the Vercel AI SDK.",
)
body("What this produced:")
bullet("Server: /api/analyze switched from generateObject to streamObject (Vercel AI SDK), returning a streamed response.")
bullet("Client: reads the stream and uses parsePartialJson to reveal each section as it parses.")
bullet(
    "UI: the audience now sees progressive status — 'Generating Summary…', "
    "'Generating Contributions…', 'Generating Methodology…' — each section "
    "filling in as Claude produces it."
)
body(
    "Trade-off: because the response is a stream (HTTP 200 once it starts), a "
    "mid-stream model error cannot change the status code — it is logged "
    "server-side and surfaced to the client via a completion check. Pre-stream "
    "errors (upload, extraction, config, rate limit) still return precise codes.",
    italic=True,
)

# ============================ STEP 9 (NEW) ============================
heading("Step 9 — Deployment  (NEW, optional)", level=1)
numbered("Push your changes to GitHub.")
numbered("Import the repository into Vercel (or your host of choice).")
numbered("Set the ANTHROPIC_API_KEY environment variable in the host's project settings.")
numbered("Deploy. The /api/analyze, /api/chat, and /api/explain routes run as server functions (Node.js runtime).")

# ============================ APPENDIX ============================
heading("Appendix — Project Structure & Troubleshooting", level=1)
body("Key files:", color=ACCENT)
bullet("app/api/analyze/route.ts — PDF upload + extraction + streamed structured analysis")
bullet("app/api/chat/route.ts — chat with the paper")
bullet("app/api/explain/route.ts — persona explanations (streaming)")
bullet("lib/pdf.ts — PDF text extraction (unpdf)")
bullet("components/upload-area.tsx — dropzone (label-based)")
bullet("components/analysis-view.tsx — progressive section cards")
bullet("components/persona-explainer.tsx — persona buttons + output")

body("Common issues:", color=ACCENT)
bullet("'Server is not configured… Set ANTHROPIC_API_KEY' — key missing; restart dev server after editing .env.local.")
bullet("'Could not extract enough text… scanned or image-only' — the PDF has no text layer (needs OCR); use a text-based PDF.")
bullet("Upload does nothing — hard-refresh (Cmd+Shift+R); if still stuck, stop the server, delete the .next folder, and re-run npm run dev.")
bullet("401 invalid x-api-key in the terminal — bad/placeholder key in .env.local.")

doc.save("Research-Paper-Explainer-Build-Guide.docx")
print("saved Research-Paper-Explainer-Build-Guide.docx")
